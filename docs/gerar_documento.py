# -*- coding: utf-8 -*-
"""
Gera o documento DOCX da atividade "Início do Projeto Mobile" para o app FinWallet.
Responde aos 4 itens exigidos: Detalhamento técnico, Contexto sociocomunitário,
Escopo do projeto e Requisitos funcionais.

Correções ABNT NBR 14724:2011 aplicadas:
  - Quebra de página embutida no parágrafo anterior (sem parágrafo isolado de quebra)
  - Nota explicativa da folha de rosto em tabela sem bordas (metade direita da folha)
  - Nota explicativa com espaço simples (1.0), conforme ABNT
  - Nomes dos autores presentes em capa E folha de rosto (obrigatório pela norma)

Saída: FinWallet_Projeto_Inicial.docx (na mesma pasta deste script)
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ----------------------------------------------------------------------------
# Dados do trabalho
# ----------------------------------------------------------------------------
INSTITUICAO = "ESTÁCIO / UNIMETA"
CURSO = "CURSO DE SISTEMAS DE INFORMAÇÃO"
DISCIPLINA = "Programação para Dispositivos Móveis em Android"
PROFESSOR = "Manah Carvalho"
ALUNOS = ["Felipe Bastos", "Pedro Barroso", "Rayton Neto", "Italo"]
CIDADE = "RIO BRANCO – ACRE"
ANO = "2026"
TITULO = "FINWALLET – APLICATIVO MOBILE DE CONTROLE FINANCEIRO PESSOAL"
SUBTITULO = "Proposta Inicial do Projeto Mobile"

BLACK = RGBColor(0x00, 0x00, 0x00)
FONT = "Arial"

# ----------------------------------------------------------------------------
# Helpers
# ----------------------------------------------------------------------------
def set_base_style(doc):
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(12)
    style.font.color.rgb = BLACK
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), FONT)
    rfonts.set(qn("w:hAnsi"), FONT)


def set_margins(doc):
    for section in doc.sections:
        section.top_margin = Cm(3)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)
        section.bottom_margin = Cm(2)


def blank(doc, n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)


def centered(doc, text, size=12, bold=False, upper=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text.upper() if upper else text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = BLACK
    return p


def page_break(doc):
    """
    Insere quebra de página EMBUTIDA no último parágrafo do documento.
    Evita criar um parágrafo isolado de quebra que alguns leitores mobile
    exibem como '--- Quebra de Página ---'.
    """
    paragraphs = doc.paragraphs
    if paragraphs:
        last_para = paragraphs[-1]
        run = last_para.add_run()
        br = OxmlElement('w:br')
        br.set(qn('w:type'), 'page')
        run._r.append(br)
    else:
        p = doc.add_paragraph()
        run = p.add_run()
        br = OxmlElement('w:br')
        br.set(qn('w:type'), 'page')
        run._r.append(br)


def _remove_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn('w:tcBorders')):
        tcPr.remove(existing)
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        el.set(qn('w:sz'), '0')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), 'auto')
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _set_table_fixed(table):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    for existing in tblPr.findall(qn('w:tblLayout')):
        tblPr.remove(existing)
    layout = OxmlElement('w:tblLayout')
    layout.set(qn('w:type'), 'fixed')
    tblPr.append(layout)
    # Largura total = área de texto (16 cm = 9072 twips)
    for existing in tblPr.findall(qn('w:tblW')):
        tblPr.remove(existing)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), '9072')
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)


def nota_explicativa(doc, text):
    """
    Nota explicativa da folha de rosto — ABNT NBR 14724:2011.
    Posicionada na metade direita da folha com espaço simples.
    Usa tabela sem bordas para garantir renderização correta em mobile.
    Coluna esquerda: espaçador (8 cm) | Coluna direita: texto (8 cm).
    """
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_fixed(table)

    left_cell, right_cell = table.rows[0].cells[0], table.rows[0].cells[1]
    _remove_cell_borders(left_cell)
    _remove_cell_borders(right_cell)

    # 8 cm + 8 cm = 16 cm (área de texto A4 com margens 3/2)
    left_cell.width = Cm(8)
    right_cell.width = Cm(8)

    p = right_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.0   # espaço simples (ABNT)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.font.name = FONT
    r.font.size = Pt(11)
    r.font.color.rgb = BLACK
    return table


def section_title(doc, number, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(f"{number} {text}")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = BLACK
    return p


def sub_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = BLACK
    return p


def body(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.color.rgb = BLACK
    return p


def bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.3
        run = p.add_run(it)
        run.font.size = Pt(12)
        run.font.color.rgb = BLACK


def two_col_table(doc, headers, rows, col0_width=None):
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = FONT
    for left, right in rows:
        cells = table.add_row().cells
        for cell, val in zip(cells, (left, right)):
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(11)
            run.font.name = FONT
    if col0_width:
        for row in table.rows:
            row.cells[0].width = col0_width
    return table


# ----------------------------------------------------------------------------
# Documento
# ----------------------------------------------------------------------------
doc = Document()
set_base_style(doc)
set_margins(doc)

# ============================== CAPA ==============================
centered(doc, INSTITUICAO, size=14, bold=True, upper=True)
centered(doc, CURSO, size=12, bold=True, upper=True)
blank(doc, 10)
centered(doc, TITULO, size=16, bold=True)
centered(doc, SUBTITULO, size=12, bold=False)
blank(doc, 14)
centered(doc, CIDADE, size=12, bold=True, upper=True)
centered(doc, ANO, size=12, bold=True)
page_break(doc)

# ======================== FOLHA DE ROSTO ========================
# Nomes dos autores ficam apenas na folha de rosto.
for aluno in ALUNOS:
    centered(doc, aluno, size=12)
blank(doc, 6)
centered(doc, TITULO, size=14, bold=True)
centered(doc, SUBTITULO, size=12)
blank(doc, 4)

nota = (
    "Trabalho apresentado à disciplina " + DISCIPLINA + ", do " + CURSO.title() +
    " da " + INSTITUICAO + ", como requisito parcial de avaliação, sob orientação "
    "do(a) Prof.(a) " + PROFESSOR + "."
)
nota_explicativa(doc, nota)   # tabela sem bordas — renderiza bem em mobile

blank(doc, 6)
centered(doc, CIDADE, size=12, bold=True, upper=True)
centered(doc, ANO, size=12, bold=True)
page_break(doc)  # quebra embutida no último parágrafo (ANO)

# ============================================================
# 1. DETALHAMENTO TÉCNICO
# ============================================================
section_title(doc, "1", "DETALHAMENTO TÉCNICO")
body(doc,
     "O FinWallet é um aplicativo mobile multiplataforma de controle financeiro "
     "pessoal, desenvolvido com foco no sistema operacional Android. A solução "
     "adota uma arquitetura offline-first: todo o processamento e o armazenamento "
     "de dados ocorrem localmente no próprio dispositivo, sem necessidade de "
     "conexão com a internet ou de servidores externos.")

sub_title(doc, "1.1 Plataforma e Linguagem")
body(doc,
     "A aplicação é executada em dispositivos móveis Android (sendo também "
     "compatível com iOS, por se tratar de uma solução multiplataforma). A "
     "linguagem de programação utilizada é o TypeScript, um superconjunto "
     "tipado do JavaScript que adiciona segurança de tipos ao desenvolvimento.")

sub_title(doc, "1.2 Framework e Ferramentas")
body(doc,
     "O desenvolvimento é baseado em React Native com a plataforma Expo, "
     "utilizando uma estrutura de componentes reutilizáveis. As principais "
     "tecnologias empregadas são apresentadas no quadro a seguir:")
two_col_table(
    doc,
    ["Tecnologia", "Descrição / Finalidade"],
    [
        ("React Native", "Framework para desenvolvimento de aplicativos móveis nativos."),
        ("Expo (SDK 54)", "Plataforma e conjunto de ferramentas para build e execução."),
        ("TypeScript", "Linguagem com tipagem estática."),
        ("Expo Router", "Navegação baseada em arquivos (file-based routing)."),
        ("SQLite (expo-sqlite)", "Banco de dados relacional local no dispositivo."),
        ("Zustand", "Gerenciamento de estado global da aplicação."),
        ("TanStack Query (React Query)", "Gerenciamento de estado assíncrono e cache de dados."),
        ("NativeWind", "Estilização com Tailwind CSS (tema claro/escuro)."),
        ("Reanimated / Moti", "Animações fluidas da interface."),
        ("Victory Native", "Gráficos (pizza e barras) para relatórios."),
        ("React Hook Form + Zod", "Construção e validação de formulários."),
        ("Expo Local Authentication", "Autenticação biométrica (digital/Face)."),
        ("Expo Notifications", "Notificações locais (lembretes e alertas)."),
        ("Expo File System + Sharing", "Exportação e compartilhamento de arquivos."),
        ("AsyncStorage", "Persistência das configurações do usuário."),
        ("Jest + Testing Library", "Testes automatizados."),
    ],
    col0_width=Cm(5.5),
)

sub_title(doc, "1.3 Arquitetura")
body(doc,
     "O sistema segue uma arquitetura em camadas com fluxo de dados "
     "unidirecional, organizada da seguinte forma: Telas (Expo Router) → "
     "Estado (Zustand) → Serviços → Banco de Dados (SQLite). Cada camada possui "
     "uma responsabilidade única:")
bullets(doc, [
    "Camada de Apresentação (app/): telas e navegação por abas e modais.",
    "Camada de Estado (stores/): armazena o estado da interface e os dados "
    "derivados (resumos e dados de gráficos) usando Zustand.",
    "Camada de Serviços (services/): concentra as regras de negócio e o acesso "
    "ao banco de dados; é a única camada que executa SQL.",
    "Camada de Dados (db/): conexão única (singleton) com o SQLite, definição "
    "do schema e carga inicial das categorias padrão.",
])
body(doc,
     "Como boas práticas de implementação, os valores monetários são "
     "armazenados em centavos (números inteiros), evitando erros de "
     "arredondamento de ponto flutuante; os identificadores dos registros são "
     "gerados como UUID; e a interface oferece tema claro e escuro. O banco de "
     "dados é composto por cinco tabelas relacionadas: categories (categorias), "
     "transactions (transações), goals (metas), goal_deposits (depósitos das "
     "metas) e budgets (orçamentos).")

sub_title(doc, "1.4 Funcionamento Geral")
body(doc,
     "Ao ser iniciado, o aplicativo inicializa o banco SQLite, cria as tabelas "
     "(caso ainda não existam) e popula as categorias padrão. Em seguida, "
     "solicita a autenticação biométrica do usuário (AuthGate). Após o acesso, "
     "é apresentada a navegação principal por abas — Início (painel), "
     "Transações, Metas e Configurações. O usuário registra receitas e despesas, "
     "define metas de economia e orçamentos por categoria, e acompanha "
     "relatórios e gráficos. As notificações locais lembram o registro diário "
     "dos gastos e avisam sobre metas próximas do vencimento.")

# ============================================================
# 2. CONTEXTO SOCIOCOMUNITÁRIO
# ============================================================
section_title(doc, "2", "CONTEXTO SOCIOCOMUNITÁRIO")

sub_title(doc, "2.1 Problema Identificado")
body(doc,
     "A baixa educação financeira é um problema social relevante no Brasil. "
     "Grande parte da população não mantém controle organizado de suas receitas "
     "e despesas, o que favorece o endividamento, a inadimplência e a "
     "dificuldade de formar reservas. A ausência de uma ferramenta simples, "
     "gratuita e acessível para acompanhar os gastos do dia a dia agrava esse "
     "cenário, especialmente entre pessoas com menor acesso à informação "
     "financeira.")

sub_title(doc, "2.2 Público-Alvo")
body(doc,
     "O aplicativo destina-se a pessoas físicas que desejam organizar suas "
     "finanças pessoais — em especial jovens, estudantes, trabalhadores "
     "autônomos e famílias de baixa e média renda — que buscam uma solução "
     "gratuita, de uso simples e que funcione mesmo sem acesso constante à "
     "internet.")

sub_title(doc, "2.3 Impacto Esperado")
body(doc,
     "Espera-se estimular a educação e o planejamento financeiro, ajudando o "
     "usuário a visualizar para onde vai o seu dinheiro, a estabelecer metas de "
     "economia e a respeitar orçamentos. Dessa forma, o projeto busca "
     "contribuir para a redução do endividamento e para a formação de hábitos "
     "financeiros mais saudáveis.")

sub_title(doc, "2.4 Benefícios para a Comunidade")
bullets(doc, [
    "Acessibilidade: funciona totalmente offline, sem consumo de dados móveis, "
    "alcançando regiões com conectividade limitada.",
    "Privacidade: os dados financeiros permanecem somente no dispositivo do "
    "usuário, sem envio a servidores externos.",
    "Gratuidade: não possui custos, democratizando o acesso ao controle "
    "financeiro.",
    "Educação financeira: relatórios e insights automáticos que orientam "
    "decisões de consumo e economia.",
    "Inclusão: interface simples, em português, com tema claro/escuro e "
    "feedback tátil (haptics).",
])

# ============================================================
# 3. ESCOPO DO PROJETO
# ============================================================
section_title(doc, "3", "ESCOPO DO PROJETO")
body(doc,
     "Esta seção delimita o escopo do projeto, descrevendo as funcionalidades "
     "contempladas (o que o sistema fará) e as exclusões assumidas nesta etapa "
     "(o que o sistema não fará).")

sub_title(doc, "3.1 O que o sistema fará")
bullets(doc, [
    "Registrar receitas e despesas com valor, descrição, categoria e data.",
    "Editar, excluir, buscar e filtrar transações.",
    "Exibir um painel (dashboard) com saldo, totais e gráficos.",
    "Permitir metas de economia com depósitos parciais e acompanhamento do "
    "progresso.",
    "Permitir orçamentos mensais por categoria, com alertas de limite.",
    "Gerar relatórios mensais, comparativos entre meses e insights automáticos.",
    "Oferecer autenticação biométrica no acesso ao aplicativo.",
    "Enviar notificações locais (lembrete diário e alerta de metas).",
    "Exportar dados em CSV e backup em JSON, além de restaurar backups.",
    "Disponibilizar tema claro e escuro.",
    "Armazenar todos os dados localmente, funcionando de forma offline.",
])

sub_title(doc, "3.2 O que o sistema não fará (limitações)")
bullets(doc, [
    "Não terá servidor/back-end nem sincronização em nuvem.",
    "Não terá conta de usuário online nem login por e-mail e senha "
    "(apenas bloqueio biométrico local).",
    "Não fará integração bancária / Open Finance, nem importação automática de "
    "extratos.",
    "Não realizará pagamentos, transferências ou qualquer movimentação "
    "financeira real.",
    "Não dará suporte a múltiplos usuários ou à sincronização entre vários "
    "dispositivos.",
    "Não fará conversão de moedas em tempo real (multimoeda).",
    "Não terá integração com órgãos públicos ou instituições financeiras.",
])

# ============================================================
# 4. REQUISITOS FUNCIONAIS
# ============================================================
section_title(doc, "4", "REQUISITOS FUNCIONAIS")
body(doc,
     "A seguir são listados os requisitos funcionais (RF) do sistema, que "
     "descrevem as funcionalidades que o aplicativo deve oferecer:")

requisitos = [
    ("RF01", "Permitir o registro de receitas e despesas, informando valor, descrição, categoria e data."),
    ("RF02", "Permitir a categorização das transações, com categorias padrão pré-cadastradas."),
    ("RF03", "Permitir editar e excluir transações."),
    ("RF04", "Permitir buscar e filtrar transações por tipo, categoria, período e descrição."),
    ("RF05", "Exibir um painel com o saldo atual, o total de receitas e o total de despesas."),
    ("RF06", "Exibir gráficos de distribuição de gastos por categoria e de evolução mensal."),
    ("RF07", "Permitir a criação de metas de economia com nome, valor-alvo e prazo."),
    ("RF08", "Permitir registrar depósitos parciais nas metas e acompanhar o progresso, concluindo a meta automaticamente ao atingir o valor-alvo."),
    ("RF09", "Permitir a criação de orçamentos mensais por categoria, com possibilidade de copiar os limites do mês anterior."),
    ("RF10", "Alertar o usuário quando os gastos atingirem 80% do orçamento ou ultrapassarem o limite definido."),
    ("RF11", "Gerar relatórios mensais com comparativo entre meses, principais categorias de gasto e insights automáticos."),
    ("RF12", "Solicitar autenticação biométrica (digital/Face ID) ao abrir o aplicativo."),
    ("RF13", "Enviar notificações locais: lembrete diário de registro e alerta de metas próximas do vencimento."),
    ("RF14", "Permitir exportar transações e metas em CSV e gerar backup completo em JSON."),
    ("RF15", "Permitir restaurar os dados a partir de um arquivo de backup JSON."),
    ("RF16", "Permitir alternar entre os temas claro e escuro."),
    ("RF17", "Armazenar todos os dados localmente, funcionando sem conexão com a internet."),
    ("RF18", "Permitir limpar todos os dados do usuário (reset), preservando as categorias padrão."),
]
two_col_table(doc, ["Código", "Descrição"], requisitos, col0_width=Cm(2.2))

# ----------------------------------------------------------------------------
# Propriedades e gravação
# ----------------------------------------------------------------------------
core = doc.core_properties
core.title = "FinWallet - Proposta Inicial do Projeto Mobile"
core.author = ", ".join(ALUNOS)
core.subject = DISCIPLINA

out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                        "FinWallet_Projeto_Inicial.docx")
doc.save(out_path)
print("DOCX gerado em:", out_path)
